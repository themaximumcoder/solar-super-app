import docx

doc = docx.Document(r'public/template.docx')

def replace_text(paragraphs):
    for p in paragraphs:
        if 'Damansara Specialist Hospital' in p.text:
            p.text = p.text.replace('Damansara Specialist Hospital', '{hospitalName} (Ph: {hospitalPhone})')
        if 'Klinik Medilove Damansara Uptown' in p.text:
            p.text = p.text.replace('Klinik Medilove Damansara Uptown', '{clinicName} (Ph: {clinicPhone})')
        if 'Damansara Fire and Rescue Station' in p.text:
            p.text = p.text.replace('Damansara Fire and Rescue Station', '{fireName} (Ph: {firePhone})')
        if 'Damansara Police Station' in p.text:
            p.text = p.text.replace('Damansara Police Station', '{policeName} (Ph: {policePhone})')

replace_text(doc.paragraphs)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            replace_text(cell.paragraphs)

doc.save(r'public/template.docx')
